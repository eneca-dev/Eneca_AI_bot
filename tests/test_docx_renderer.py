"""Tests for docx_renderer.

Verifies that:
  * the output is a valid .docx (bytes parseable by python-docx)
  * key text content is present
  * the transcript document includes the 3-column table
"""
from io import BytesIO

import pytest
from docx import Document

from agents.teams_agent import (
    Author, DiscussionItem, MeetingParticipant, MeetingReport,
    MeetingTranscript, OpenQuestion, Risk, TranscriptSegment,
)
from services.docx_renderer import render_report_docx, render_transcript_docx


# --- Fixtures ---


def _sample_report() -> MeetingReport:
    return MeetingReport(
        subject="Совещание по фундаментам",
        date="2026-04-27",
        duration="1h 5m",
        location="Microsoft Teams",
        project="ЖК Северный",
        transcript_url=None,
        previous_protocol_url=None,
        participants=[
            MeetingParticipant(organization="Eneca", name="Иван Петров", role="Инженер"),
            MeetingParticipant(organization=None, name="Мария Сидорова", role=None),
        ],
        preview_summary="Решено использовать буронабивные сваи.",
        discussion_items=[
            DiscussionItem(topic="Тип свай", outcome="Решено: буронабивные",
                           responsible="Иван", deadline="пятница", status="Новый"),
        ],
        open_questions=[
            OpenQuestion(question="Готовы ли испытания грунта?", responsible="Мария",
                         deadline="среда", comment="Перед началом работ"),
        ],
        risks=[],
        author=Author(organization="Eneca", name="Иван Петров", role="Инженер"),
    )


def _sample_transcript() -> MeetingTranscript:
    return MeetingTranscript(
        title="Совещание по фундаментам",
        date="2026-04-27",
        duration="1h 5m",
        participants=[
            MeetingParticipant(organization="Eneca", name="Иван Петров", role="Инженер"),
        ],
        transcript=[
            TranscriptSegment(speaker="Иван Петров", timestamp="00:00:15",
                              text="Коллеги, начнём."),
            TranscriptSegment(speaker="Иван Петров", timestamp="00:01:00",
                              text="Решаем по сваям."),
        ],
    )


# --- Report ---


def test_report_docx_returns_bytes():
    out = render_report_docx(_sample_report())
    assert isinstance(out, (bytes, bytearray))
    assert len(out) > 0


def test_report_docx_is_parseable():
    out = render_report_docx(_sample_report())
    doc = Document(BytesIO(out))
    # Should have paragraphs
    assert len(doc.paragraphs) > 0


def test_report_docx_contains_header_fields():
    out = render_report_docx(_sample_report())
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Протокол совещания от 2026-04-27" in text
    assert "Microsoft Teams" in text
    assert "ЖК Северный" in text
    assert "Совещание по фундаментам" in text
    assert "Иван Петров" in text


def test_report_docx_contains_section_headings():
    out = render_report_docx(_sample_report())
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "1. Обсужденные вопросы" in text
    assert "2. Открытые вопросы" in text
    assert "3. Риски" in text


def test_report_docx_empty_section_shows_placeholder():
    out = render_report_docx(_sample_report())
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    # Risks list is empty → "Нет"
    risks_idx = text.index("3. Риски")
    assert "Нет" in text[risks_idx:risks_idx + 200]


def test_report_docx_author_line():
    out = render_report_docx(_sample_report())
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Составитель: Eneca, Иван Петров, Инженер" in text


def test_report_docx_dash_for_empty_optional_fields():
    report = _sample_report()
    report.project = None
    out = render_report_docx(report)
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Проект: —" in text


# --- Transcript ---


def test_transcript_docx_returns_bytes():
    out = render_transcript_docx(_sample_transcript())
    assert isinstance(out, (bytes, bytearray))
    assert len(out) > 0


def test_transcript_docx_includes_three_column_table():
    out = render_transcript_docx(_sample_transcript())
    doc = Document(BytesIO(out))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # 1 header row + 2 segment rows
    assert len(table.rows) == 3
    assert len(table.columns) == 3
    head_cells = [c.text for c in table.rows[0].cells]
    assert head_cells == ["Время", "Спикер", "Реплика"]


def test_transcript_docx_table_contains_segment_data():
    out = render_transcript_docx(_sample_transcript())
    doc = Document(BytesIO(out))
    table = doc.tables[0]
    row1 = [c.text for c in table.rows[1].cells]
    assert row1 == ["00:00:15", "Иван Петров", "Коллеги, начнём."]


def test_transcript_docx_handles_empty_transcript():
    transcript = _sample_transcript()
    transcript.transcript = []
    out = render_transcript_docx(transcript)
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Нет" in text
    assert len(doc.tables) == 0
