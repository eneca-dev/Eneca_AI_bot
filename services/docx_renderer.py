"""DOCX renderers for the meeting protocol and the transcript.

Two outputs:
  * `render_report_docx(report)` — the same structure as the Teams text
    message, but as a Word document. Headings + bulleted/numbered sections.
  * `render_transcript_docx(transcript)` — header + a 3-column table
    (Время / Спикер / Реплика).

Both functions return raw bytes (no temp file is written) so the caller
can stream them straight into Supabase Storage.
"""
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.teams_agent import MeetingReport, MeetingTranscript


_DASH = "—"
_EMPTY_SECTION = "Нет"


def _val(v) -> str:
    if v is None:
        return _DASH
    s = str(v).strip()
    return s if s else _DASH


# ---------- Report (protocol) ----------

def _add_kv(doc, key: str, value: str) -> None:
    """Bold key on one line followed by the value."""
    p = doc.add_paragraph()
    run = p.add_run(f"{key}: ")
    run.bold = True
    p.add_run(value)


def _add_section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _render_discussion_section(doc, items) -> None:
    _add_section_heading(doc, "1. Обсужденные вопросы, договоренности и действия")
    if not items:
        doc.add_paragraph(_EMPTY_SECTION)
        return
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. Вопрос/тема: ").bold = True
        p.add_run(_val(item.topic))
        doc.add_paragraph(f"   Итог/действие: {_val(item.outcome)}")
        doc.add_paragraph(f"   Ответственный: {_val(item.responsible)}")
        doc.add_paragraph(f"   Срок: {_val(item.deadline)}")
        doc.add_paragraph(f"   Статус: {_val(item.status)}")


def _render_open_questions_section(doc, items) -> None:
    _add_section_heading(doc, "2. Открытые вопросы")
    if not items:
        doc.add_paragraph(_EMPTY_SECTION)
        return
    for i, q in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. Вопрос: ").bold = True
        p.add_run(_val(q.question))
        doc.add_paragraph(f"   Ответственный: {_val(q.responsible)}")
        doc.add_paragraph(f"   Срок получения ответа: {_val(q.deadline)}")
        doc.add_paragraph(f"   Комментарий: {_val(q.comment)}")


def _render_risks_section(doc, items) -> None:
    _add_section_heading(doc, "3. Риски")
    if not items:
        doc.add_paragraph(_EMPTY_SECTION)
        return
    for i, r in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. Риск: ").bold = True
        p.add_run(_val(r.risk))
        doc.add_paragraph(f"   Причина: {_val(r.cause)}")
        doc.add_paragraph(f"   Возможные последствия: {_val(r.consequences)}")
        doc.add_paragraph(f"   Ответственный: {_val(r.responsible)}")
        doc.add_paragraph(f"   Действие: {_val(r.mitigation)}")


def render_report_docx(report: MeetingReport) -> bytes:
    """Render a MeetingReport as a .docx and return the bytes."""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Протокол совещания от {_val(report.date)}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Header block
    _add_kv(doc, "Место проведения", _val(report.location))
    _add_kv(doc, "Дата", _val(report.date))
    _add_kv(doc, "Проект", _val(report.project))
    _add_kv(doc, "Транскрибация", _val(report.transcript_url))
    _add_kv(doc, "Предмет совещания", _val(report.subject))
    _add_kv(doc, "Ссылка на протокол предшествующего совещания", _val(report.previous_protocol_url))

    # Participants
    _add_section_heading(doc, "Участники")
    if report.participants:
        for i, p in enumerate(report.participants, 1):
            doc.add_paragraph(
                f"{i}. Организация: {_val(p.organization)} | "
                f"ФИО: {_val(p.name)} | "
                f"Должность: {_val(p.role)}"
            )
    else:
        doc.add_paragraph(_EMPTY_SECTION)

    # Preview summary (only if present)
    if report.preview_summary:
        p = doc.add_paragraph()
        p.add_run("Резюме: ").bold = True
        p.add_run(report.preview_summary)

    # Three template sections
    _render_discussion_section(doc, report.discussion_items)
    _render_open_questions_section(doc, report.open_questions)
    _render_risks_section(doc, report.risks)

    # Author
    author = report.author
    parts = [v for v in (author.organization, author.name, author.role) if v]
    p = doc.add_paragraph()
    p.add_run("Составитель: ").bold = True
    p.add_run(", ".join(parts) if parts else _DASH)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------- Transcript ----------

def render_transcript_docx(transcript: MeetingTranscript) -> bytes:
    """Render a MeetingTranscript as .docx — header + 3-column table."""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Транскрипт встречи: {_val(transcript.title)}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Meta
    _add_kv(doc, "Дата", _val(transcript.date))
    _add_kv(doc, "Длительность", _val(transcript.duration))

    # Participants
    if transcript.participants:
        _add_section_heading(doc, "Участники")
        for i, p in enumerate(transcript.participants, 1):
            doc.add_paragraph(
                f"{i}. Организация: {_val(p.organization)} | "
                f"ФИО: {_val(p.name)} | "
                f"Должность: {_val(p.role)}"
            )

    # Transcript table
    _add_section_heading(doc, "Транскрипт")
    if not transcript.transcript:
        doc.add_paragraph(_EMPTY_SECTION)
    else:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        head = table.rows[0].cells
        head[0].text = "Время"
        head[1].text = "Спикер"
        head[2].text = "Реплика"
        for cell in head:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for seg in transcript.transcript:
            row = table.add_row().cells
            row[0].text = _val(seg.timestamp)
            row[1].text = _val(seg.speaker)
            row[2].text = _val(seg.text)

        # Reasonable column widths (approximate; Word may auto-adjust)
        for row in table.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(3.5)
            row.cells[2].width = Cm(11.0)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
